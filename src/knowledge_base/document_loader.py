import pdfplumber
import re
import os
import tempfile
from typing import List

class DocumentLoader:
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.doc', '.docx']

    def load_file(self, file_path: str) -> str:
        """Load content from a file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.supported_formats}")

        content = ""
        try:
            if ext == '.pdf':
                content = self._load_pdf(file_path)
            elif ext == '.txt':
                content = self._load_txt(file_path)
            elif ext == '.docx':
                content = self._load_docx(file_path)
            elif ext == '.doc':
                content = self._load_doc(file_path)
        except Exception as e:
            raise RuntimeError(f"Error loading file {file_path}: {str(e)}")
        
        return self.clean_text(content)

    def _load_pdf(self, file_path: str) -> str:
        content = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        return content

    def _load_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _load_docx(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "DOCX support requires the 'python-docx' package to be installed."
            ) from exc

        document = Document(file_path)
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _load_doc(self, file_path: str) -> str:
        if os.name != "nt":
            raise RuntimeError("DOC support currently requires Windows with Microsoft Word installed.")

        temp_txt_path = None
        word_app = None
        document = None

        try:
            try:
                import win32com.client
            except ImportError as exc:
                raise RuntimeError(
                    "DOC support requires the 'pywin32' package and Microsoft Word on Windows."
                ) from exc

            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
            temp_txt_path = temp_file.name
            temp_file.close()

            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            document = word_app.Documents.Open(os.path.abspath(file_path))
            document.SaveAs(os.path.abspath(temp_txt_path), FileFormat=2)
            document.Close(False)
            document = None
            word_app.Quit()
            word_app = None

            return self._load_txt(temp_txt_path)
        except Exception as exc:
            raise RuntimeError(
                "DOC support failed. Ensure Microsoft Word is installed and accessible on this machine."
            ) from exc
        finally:
            if document is not None:
                document.Close(False)
            if word_app is not None:
                word_app.Quit()
            if temp_txt_path and os.path.exists(temp_txt_path):
                os.remove(temp_txt_path)

    def clean_text(self, text: str) -> str:
        """Clean noise from financial reports."""
        if not text:
            return ""
            
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # 1. Strip whitespace
            line = line.strip()
            
            # 2. Skip empty lines
            if not line:
                continue
                
            # 3. Skip page numbers (e.g., "1", "Page 1", "1 / 10")
            if re.match(r'^page\s*\d+(\s*[/of-]\s*\d+)?$', line, re.IGNORECASE) or re.match(r'^\d+$', line):
                continue
                
            # 4. Skip common disclaimer headers (simplified)
            if "disclaimer" in line.lower() and len(line) < 50:
                continue
                
            cleaned_lines.append(line)
        
        return "\n".join(cleaned_lines)
